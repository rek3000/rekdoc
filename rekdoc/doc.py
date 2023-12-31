#!/usr/bin/env python
###
import sys
import os
import logging
import json
import docx

#
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

#
from rekdoc import tools

TABLE_RED = "#C00000"
ASSERTION = {1: "Kém", 3: "Cần lưu ý", 5: "Tốt"}


def assert_fault(data):
    if data["fault"] == "":
        score = ""
        comment = [""]
    elif data["fault"] == "No faults found":
        score = 5
        comment = ["Lỗi: Không", "Đánh giá: " + ASSERTION[score]]
    elif data["fault"] == "warning":
        score = 3
        comment = [
            "Lỗi không ảnh hưởng tới hiệu năng của hệ thống",
            "Đánh giá: " + ASSERTION[score],
        ]
    else:
        score = 1
        comment = [
            "Lỗi ảnh hưởng tới hoạt động của hệ thống",
            "Đánh giá: " + ASSERTION[score],
        ]

    fault = [score, comment]
    logging.debug(json.dumps(fault, ensure_ascii=False))
    return fault


def assert_temp(data):
    inlet_temp = data["inlet"].split()[0]
    inlet_temp = int(inlet_temp)
    if inlet_temp == "":
        score = ""
        comment = [""]
    elif inlet_temp >= 21 and inlet_temp <= 23:
        score = 5
        comment = [
            "Nhiệt độ bên trong: " + str(inlet_temp),
            "Đánh giá: " + ASSERTION[score],
        ]
    elif inlet_temp > 23 and inlet_temp <= 26:
        score = 3
        comment = [
            "Nhiệt độ bên trong: " + str(inlet_temp),
            "Đánh giá: " + ASSERTION[score],
        ]
    elif inlet_temp > 26:
        score = 1
        comment = [
            "Nhiệt độ bên trong: " + str(inlet_temp),
            "Đánh giá: " + ASSERTION[score],
        ]

    temp = [score, comment]
    logging.debug(json.dumps(temp, ensure_ascii=False))
    return temp


def assert_firmware(data):
    latest = ""
    if data["firmware"] == "":
        firmware = ["", [""]]
        logging.debug(json.dumps(firmware, ensure_ascii=False))
        return firmware
    while True:
        try:
            sys.stdout.write("\033[?25h")
            latest = (
                input(
                    "Enter latest ILOM version\n[" + data["firmware"] + "] "
                )
                or data["firmware"]
            )
        except KeyboardInterrupt:
            print()
            sys.exit()
        except ValueError:
            continue
        break

    if latest == data["firmware"]:
        score = 5
    else:
        while True:
            try:
                print("Đánh giá")
                print("[0] Tốt")
                print("[1] Cần lưu ý")
                print("[2] Kém")
                score = int(input("Chọn đánh giá\n [0] ") or "0")
            except KeyboardInterrupt:
                print()
                sys.exit()
            except ValueError:
                continue
            break

    comment = [
        "Phiên bản ILOM hiện tại: " + data["firmware"],
        "Phiên bản ILOM mới nhất: " + latest,
        "Đánh giá: " + ASSERTION[score]
    ]
    firmware = [score, comment]
    logging.debug(json.dumps(firmware, ensure_ascii=False))
    return firmware


def assert_image(data):
    score = 0
    if data["image"] == "":
        image = ["", [""]]
        logging.debug(json.dumps(image, ensure_ascii=False))
        return image

    while True:
        try:
            sys.stdout.write("\033[?25h")
            latest = (
                input("Enter latest OS version\n[" + data["image"] + "] ")
                or data["image"]
            )
        except KeyboardInterrupt:
            print()
            sys.exit()
        except ValueError:
            continue
        break

    if latest == data["image"]:
        score = 5
    else:
        while True:
            try:
                print("Đánh giá")
                print("[0] Tốt")
                print("[1] Cần lưu ý")
                print("[2] Kém")
                score = int(input("Chọn đánh giá\n [0] ") or "0")
            except KeyboardInterrupt:
                print()
                sys.exit()
            except ValueError:
                continue
            break

    comment = [
        "Phiên bản OS hiện tại: " + data["image"],
        "Phiên bản OS mới nhất: " + latest,
        "Đánh giá: " + ASSERTION[score]
    ]
    image = [score, comment]
    logging.debug(json.dumps(image, ensure_ascii=False))
    return image


def assert_vol(data):
    score = 0
    if data["vol_avail"] == "":
        score = ""
        comment = [""]
    elif data["vol_avail"] > 30 and data["raid_stat"] is True:
        score = 5
        comment = [
            "Phân vùng OS được cấu hình RAID",
            "Dung lượng khả dụng: " + str(data["vol_avail"]) + "%",
            "Đánh giá: " + ASSERTION[score],
        ]
    elif (data["vol_avail"] > 15 and data["vol_avail"] <= 30) and data[
        "raid_stat"
    ] is True:
        score = 3
        comment = [
            "Phân vùng OS được cấu hình RAID",
            "Dung lượng khả dụng: " + str(data["vol_avail"]) + "%",
            "Đánh giá: " + ASSERTION[score],
        ]
    elif data["vol_avail"] <= 15 and data["raid_stat"] is False:
        score = 1
        comment = [
            "Phân vùng OS không được cấu hình RAID",
            "Dung lượng khả dụng: " + str(data["vol_avail"]) + "%",
            "Đánh giá: " + ASSERTION[score],
        ]
    elif data["vol_avail"] <= 15 and data["raid_stat"] is True:
        score = 1
        comment = [
            "Phân vùng OS được cấu hình RAID",
            "Dung lượng khả dụng: " + str(data["vol_avail"]) + "%",
            "Đánh giá: " + ASSERTION[score],
        ]
    elif data["vol_avail"] > 30 and data["raid_stat"] is False:
        score = 3
        comment = [
            "Phân vùng OS không được cấu hình RAID",
            "Dung lượng khả dụng: " + str(data["vol_avail"]) + "%",
            "Đánh giá: " + ASSERTION[score],
        ]

    vol = [score, comment]
    logging.debug(json.dumps(vol, ensure_ascii=False))

    return vol


def assert_bonding(data):
    if data["bonding"] == "":
        score = ""
        comment = [""]
    elif data["bonding"] == "none":
        score = 1
        comment = ["Network không được cấu hình bonding"]
    elif data["bonding"] == "aggr":
        score = 5
        comment = ["Network được cấu hình bonding Aggregration"]
    elif data["bonding"] == "ipmp":
        score = 5
        comment = ["Network được cấu hình bonding IPMP"]
    comment.append("Đánh giá: " + ASSERTION[score])

    bonding = [score, comment]
    logging.debug(json.dumps(bonding, ensure_ascii=False))

    return bonding


def assert_cpu_util(data):
    if data["cpu_util"] == "":
        score = ""
        comment = [""]
    elif data["cpu_util"] <= 30:
        score = 5
    elif data["cpu_util"] > 30 and data["cpu_util"] <= 70:
        score = 3
    else:
        score = 1
    comment = ["CPU Utilization khoảng " + str(data["cpu_util"]) + "%"]
    comment.append("Đánh giá: " + ASSERTION[score])

    cpu_util = [score, comment]
    logging.debug(json.dumps(cpu_util, ensure_ascii=False))

    return cpu_util


def assert_load(data):
    if data["load_avg"] == "":
        score = ""
        comment = [""]
    elif data["load"]["load_avg_per"] <= 2:
        score = 5
    elif 2 < data["load"]["load_avg_per"] <= 5:
        score = 3
    else:
        score = 1
    comment = [
        "CPU Load Average: " + str(data["load"]["load_avg"]),
        "Number of Cores: " + str(data["load"]["vcpu"]),
        "CPU Load Average per core = CPU Load Average / Number of Cores = "
        + str(data["load"]["load_avg_per"]),
        "Đánh giá: " + ASSERTION[score]
    ]

    load = [score, comment]
    logging.debug(json.dumps(load, ensure_ascii=False))

    return load


def assert_mem_free(data):
    # mem_free = 100 - data["mem_util"]
    mem_free = data["mem_free"]
    if mem_free == "":
        score = ""
        comment = [""]
        mem_free = [score, comment]
        logging.debug(json.dumps(mem_free, ensure_ascii=False))
        return mem_free
    elif mem_free >= 20:
        score = 5
    elif mem_free > 10 and mem_free < 20:
        score = 3
    else:
        score = 1
    comment = ["Average physical memory free: " + str(mem_free) + "%"]
    comment.append("Đánh giá: " + ASSERTION[score])

    mem_free = [score, comment]
    logging.debug(json.dumps(mem_free, ensure_ascii=False))

    return mem_free


def assert_io_busy(data):
    if data["io_busy"]["busy"] < 50:
        score = 5
    elif 50 <= data["io_busy"]["busy"] <= 70:
        score = 3
    else:
        score = 1

    comment = ["Thiết bị IO Busy cao: " + data["io_busy"]["name"]]
    comment.append("IO Busy: " + str(data["io_busy"]["busy"]))
    comment.append("Đánh giá: " + ASSERTION[score])
    io_busy = [score, comment]
    return io_busy


def assert_swap_util(data):
    if data["swap_util"] <= 2:
        score = 5
    elif data["swap_util"] > 2 and data["swap_util"] <= 5:
        score = 3
    else:
        score = 1
    comment = ["SWAP Utilization: " + str(data["swap_util"]) + "%"]

    swap_util = [score, comment]
    logging.debug(json.dumps(swap_util, ensure_ascii=False))

    return swap_util


def assert_ilom(data):
    x = {}
    try:
        fault = assert_fault(data)
        temp = assert_temp(data)
        firmware = assert_firmware(data)
        x = {"fault": fault,
             "temp": temp,
             "firmware": firmware}
    except RuntimeError:
        print("Failed to assert ILOM")
        raise
    return x


def assert_system_status(data, server_type):
    x = {}
    try:
        image = assert_image(data)
        vol = assert_vol(data)
        bonding = assert_bonding(data)
        x = {"image": image,
             "vol": vol,
             "bonding": bonding}
    except RuntimeError:
        print("Failed to assert system status")
        raise
    return x


def assert_system_perform(data, platform, system_type):
    x = {}
    try:
        if system_type == "standalone":
            if platform == "solaris":
                cpu_util = assert_cpu_util(data)
                mem_free = assert_mem_free(data)
                io_busy = assert_io_busy(data)

                x = {"cpu_util": cpu_util,
                     "mem_free": mem_free,
                     "io_busy": io_busy}
            elif platform == "linux":
                pass
        elif system_type == "exa":
            pass
    except RuntimeError:
        print("Failed to assert system performance")
        raise
    logging.debug(json.dumps(x))
    return x


def assert_data(data):
    asserted = {}
    # for i in data:
    #     if i == "inlet":
    #         i = "temp"
    #     if i == "exhaust":
    #         continue
    #     if i == "raid_stat":
    #         continue
    #     # if i == "mem_util":
    #     #     i = "mem_free"
    #     asserted[i] = ["", []]

    ilom = assert_ilom(data)
    system_status = None
    system_perform = None
    if system_info["system_type"] == "standalone":
        system_status = assert_system_status(data,
                                             # system_info["platform"],
                                             system_info["type"])
        system_perform = assert_system_perform(data,
                                               system_info["platform"],
                                               system_info["system_type"])
    elif system_info["system_type"] == "exa":
        system_status = assert_system_status(data,
                                             # system_info["system_type"],
                                             system_info["type"])
        system_perform = assert_system_perform(data,
                                               system_info["platform"],
                                               system_info["system_type"])
    asserted = {"node_name": data["node_name"],
                **ilom,
                **system_status,
                **system_perform}

    # cpu_util = assert_cpu_util(data)
    # asserted["cpu_util"][0] = cpu_util[0]
    # asserted["cpu_util"][1].extend(cpu_util[1])
    #
    # mem_free = assert_mem_free(data)
    # asserted["mem_free"][0] = mem_free[0]
    # asserted["mem_free"][1].extend(mem_free[1])
    # load = assert_load(data)
    # asserted["load"][0] = load[0]
    # asserted["load"][1].extend(load[1])
    #
    #
    # swap_util = assert_swap_util(data)
    # asserted["swap_util"][0] = swap_util[0]
    # asserted["swap_util"][1].extend(swap_util[1])

    for field in asserted:
        logging.debug("ASSERTED:" + field + ": " + str(asserted[field][0]))
    return asserted


def get_score(asserted):
    checklist = [
        ["STT", "Hạng mục kiểm tra", "Score"],
        [1, "Kiểm tra trạng thái phần cứng", ["", []]],
        [2, "Kiểm tra nhiệt độ", ["", []]],
        [3, "Kiểm tra phiên bản ILOM", ["", []]],
        [4, "Kiểm tra phiên bản Image", ["", []]],
        [5, "Kiểm tra cấu hình RAID và dung lượng phân vùng OS", ["", []]],
        [6, "Kiểm tra cấu hình Bonding Network", ["", []]],
        [7, "Kiểm tra CPU Utilization", ["", []]],
        # [8, "Kiểm tra CPU Load Average", ["", []]],
        [8, "Kiểm tra Memory", ["", []]],
        [9, "Kiểm tra IO Busy", ["", []]],
    ]

    keys = list(asserted)

    for i in range(1, len(checklist)):
        asserted_score = asserted[keys[i]][0]
        comment = asserted[keys[i]][1]
        try:
            score = ASSERTION[asserted_score]
        except Exception:
            score = asserted_score
        checklist[i][2][0] = score
        logging.info(checklist[i][1] + ":" + score)
        checklist[i][2][1] = comment

    return checklist


# This function table with last column cells may or may not contain
# list of string
def drw_table(doc, checklist, row, col, info=False):
    if checklist == []:
        return -1
    tab = doc.add_table(row, col)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.style = "Table Grid"

    # ADD TITLE CELLS AND COLOR THEM
    cols = tab.rows[0].cells
    for r in range(len(checklist[0])):
        cell = cols[r]
        cell.text = checklist[0][r]
        cell.paragraphs[0].style = "Table Heading"
        shading_elm = parse_xml(
            r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), TABLE_RED)
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # ADD CONTENT TO TABLE
    if not info:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                cells[j].text = checklist[i][j]
                cells[j].paragraphs[0].style = "Table Paragraph"
    else:
        for i in range(1, len(checklist)):
            rows = tab.rows[i]
            cells = rows.cells
            for j in range(0, len(checklist[i])):
                if j == (len(checklist[i]) - 1):
                    cells[j].text = checklist[i][j][0]
                    cells[j].paragraphs[0].style = "Table Paragraph"
                    continue
                cells[j].text = str(checklist[i][j])
                cells[j].paragraphs[0].style = "Table Paragraph"
    return tab


# def drw_image_to_doc(doc, node, images_root, images_name):
# path?
def drw_info(doc, node, checklist, images_root, images_name=[]):
    for i in range(1, len(checklist)):
        doc.add_paragraph(checklist[i][1], style="baocao4")
        try:
            if isinstance(images_name[i - 1], list):
                for image in images_name[i - 1]:
                    path = os.path.normpath(
                        images_root + "/" + node + "/" + image)
                    doc.add_picture(path, width=Inches(6.73))
            else:
                path = os.path.normpath(
                    images_root + "/" + node + "/" + images_name[i - 1])
                doc.add_picture(
                    path,
                    width=Inches(6.73),
                )
        except Exception:
            pass
        for line in checklist[i][2][1]:
            doc.add_paragraph(line, style="Dash List")
    doc.add_page_break()


def define_doc(sample):
    try:
        doc = docx.Document(os.path.normpath(sample))
    except Exception:
        print("Sample docx not found!")
        sys.exit()
    return doc


def drw_menu(doc, nodes):
    # doc.add_paragraph("ORACLE EXADATA X8M-2", style="baocao1")
    doc.add_paragraph("Kiểm tra nhiệt độ môi trường", style="baocao2")
    doc.add_paragraph("Mục lục", style="Heading")
    for node in nodes:
        doc.add_paragraph("Kiểm tra nhiệt độ môi trường", style="baocao2")
        doc.add_paragraph("").paragraph_format.tab_stops.add_tab_stop(
            Inches(1.5), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS
        )
    doc.add_page_break()


system_info = {"system_type": "",
               "platform": "",
               "type": ""}


def drw_doc(doc, input_file, out_dir, images_root, force):
    input_file_data = tools.read_json(input_file)
    system_info["system_type"] = input_file_data["system_type"]
    system_info["platform"] = input_file_data["platform"]
    system_info["type"] = input_file_data["type"]
    nodes = input_file_data["nodes"]
    if nodes == -1:
        return -1
    asserted_list = []
    doc.add_page_break()
    # drw_menu(doc, nodes)
    input_root = os.path.split(input_file)[0]
    for node in nodes:
        print("NODE:" + node["node_name"])
        print("RUNNING:GETTING SAVED IMAGES")
        image_json = os.path.normpath(
            images_root + "/" + node["node_name"] + "/images.json")
        images_name = tools.read_json(image_json)

        print("RUNNING:ASSERTING DATA")
        asserted = assert_data(node)

        print("RUNNING:SAVING ASSERTED DATA")
        file_dump = asserted
        asserted_file = input_root + "/" + \
            node["node_name"] + "/" + node["node_name"] + "_asserted.json"
        asserted_list += [asserted_file]
        tools.save_json(
            os.path.normpath(asserted_file),
            file_dump,
        )
        print("RUNNING:CREATING CHECKLIST")
        checklist = get_score(asserted)
        print("RUNNING:DRAWING OVERVIEW TABLE")
        doc.add_paragraph("Máy chủ " + node["node_name"], style="baocao2")
        doc.add_paragraph("Thông tin tổng quát", style="baocao3")
        overview = [
            ["Hostname", "Product Name", "Serial Number", "IP Address"],
            [node["node_name"], "", "", ""],
        ]
        drw_table(doc, overview, 2, 4)
        doc.add_paragraph("Đánh giá", style="baocao3")
        print("RUNNING:DRAWING SUMMARY TABLE")
        drw_table(doc, checklist, len(checklist), 3, True)

        print("RUNNING:DRAWING DETAILS")
        doc.add_paragraph("Thông tin chi tiết", style="baocao3")
        drw_info(doc, node["node_name"], checklist, images_root, images_name)

        print("DONE")
        print()
    logging.debug(json.dumps(asserted_list))
    print("RUNNING:SAVING ASSERTED SUMMARY FILE")
    file_name = os.path.normpath(tools.rm_ext(
        input_file, "json") + "_asserted.json")
    tools.join_json(file_name, asserted_list)
    return doc


def print_style(doc):
    styles = doc.styles
    for style in styles:
        print(style.name)


def run(input_file, output_file, sample, images_dir, force=False):
    doc = define_doc(sample)
    out_dir = os.path.split(output_file)[0]
    try:
        doc = drw_doc(doc, input_file, out_dir, images_dir, force)
        if doc == -1:
            return -1
    except Exception as err:
        print(err)
        return -1

    if logging.root.level == 10:
        print()
        print("List of all styles")
        print_style(doc)
        print()
    file_name = os.path.normpath(tools.rm_ext(output_file, "json") + ".docx")
    doc.save(file_name)
    return file_name


##### MAIN #####
def main():
    run()


if __name__ == "__main__":
    main()
